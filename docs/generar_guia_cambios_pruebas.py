from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "docs/guia-cambios-y-pruebas-anaya-2026-06-26.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths):
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)
        set_cell_shading(header_cells[index], "E8EEF5")

    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            set_cell_text(cells[index], value)

    if widths:
        set_table_widths(table, widths)

    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    return paragraph


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Guia De Cambios Y Verificacion Funcional")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Sistema Finca Anaya - Cambios trabajados el 26 de junio de 2026").bold = True
    subtitle.paragraph_format.space_after = Pt(12)

    intro = doc.add_paragraph()
    intro.add_run(
        "Este documento resume los cambios aplicados durante el dia y entrega una guia paso a paso "
        "para validar las funcionalidades principales con datos reales o de prueba. La intencion es "
        "que la persona enlace de la empresa pueda revisar el flujo completo, registrar hallazgos y "
        "confirmar si el sistema se ajusta a la operacion diaria."
    )


def add_summary(doc):
    doc.add_heading("1. Resumen Ejecutivo", level=1)
    add_bullet(doc, "Se reforzo el flujo desde cotizacion hasta despacho, separando mejor responsabilidades por rol.")
    add_bullet(doc, "Bodega ahora tiene mejor control de pedidos pendientes, prioridad, alistamiento y despacho.")
    add_bullet(doc, "Laboratorio controla el inicio de procesos, la finalizacion fisica, el examen final y la orden de mezcla.")
    add_bullet(doc, "El cafe procesado no vuelve a disponibilidad operativa hasta que laboratorio registre los examenes finales.")
    add_bullet(doc, "Se corrigio el manejo de abonos multiples y el historial de pagos en documentos.")
    add_bullet(doc, "Se mejoraron filtros, dashboard por rol y navegabilidad sin cambiar los nombres de modulos usados por la empresa.")


def add_changes(doc):
    doc.add_heading("2. Cambios Implementados Durante El Dia", level=1)
    rows = [
        ("Flujo comercial", "El vendedor cotiza perfiles/tipos de cafe sin seleccionar lotes.", "Reduce errores y deja la decision operativa a bodega."),
        ("Ventas aprobadas", "Al convertir cotizacion en venta, queda pendiente para bodega.", "Bodega recibe una lista clara de pedidos por alistar."),
        ("Pendientes de bodega", "Se mejoro la vista con prioridad, fecha, estado y siguiente accion.", "Permite ordenar el trabajo diario por urgencia."),
        ("Prioridad de entrega", "Bodega puede clasificar prioridad alta, media o baja.", "Ayuda a decidir que pedido se atiende primero."),
        ("Solicitud de proceso", "Crear proceso ya no descuenta inventario inmediatamente.", "Evita restar cafe antes de confirmar que entro a proceso."),
        ("Inicio de proceso", "Laboratorio confirma inicio, ubicacion y fecha estimada de regreso.", "Ahi se descuentan los lotes usados y queda trazabilidad."),
        ("Cafe en proceso", "El sistema muestra que el cafe esta en proceso y no disponible.", "Evita vender o alistar cafe que todavia no esta listo."),
        ("Proceso terminado", "Se agrego paso de proceso fisico terminado y pendiente de laboratorio.", "El cafe no pasa directo a bodega sin examen."),
        ("Examen final", "Laboratorio registra humedad final, catacion, score, peso final y perfil.", "Solo despues se crea el lote PROC disponible."),
        ("Lote procesado", "Al finalizar laboratorio se genera un lote PROC con datos tecnicos.", "Permite vender o alistar cafe procesado con informacion completa."),
        ("Mezcla final", "Laboratorio define porcentajes por lote/categoria para la venta.", "Bodega recibe una orden clara para ensamblar el perfil solicitado."),
        ("Orden para bodega", "La orden impresa muestra pedido, lotes, porcentajes y kg estimados.", "Puede imprimirse o guardarse como PDF para entregar a quien mezcla."),
        ("Descuento de inventario", "Venta directa descuenta al marcar alistada; proceso descuenta al iniciar.", "Evita dobles descuentos y mantiene inventario coherente."),
        ("Despacho", "Despachar no vuelve a descontar inventario.", "El descuento ya ocurrio en alistamiento o inicio de proceso."),
        ("Abonos", "Se corrigio el registro de varios pagos sobre una misma venta.", "El sistema recalcula pagado, saldo y estado."),
        ("Pago completo", "Cuando el saldo llega a cero, la venta pasa a pagada.", "Contabilidad ve cartera actualizada automaticamente."),
        ("Documentos", "La factura/documento muestra historial de pagos.", "Sustenta pagos con fecha, metodo, referencia, valor y notas."),
        ("Logs backend", "Se agregaron logs de depuracion para pagos y errores.", "Facilita detectar problemas en Render o consola."),
        ("Comercial", "Se agregaron filtros por estado y cliente rapido en cotizacion.", "El vendedor no debe salir del flujo para crear un cliente."),
        ("Ventas", "Se agregaron filtros operativos y de pago, con textos por rol.", "Cada usuario ve la pantalla con enfoque mas claro."),
        ("Muestras", "Se agregaron filtros y orden por estado.", "La persona encargada ve primero lo pendiente."),
        ("Dashboard", "Se ordenaron tarjetas segun rol.", "Cada usuario ve primero lo que le corresponde revisar."),
    ]
    add_table(doc, ["Area", "Cambio", "Resultado esperado"], rows, [Inches(1.35), Inches(3.0), Inches(2.4)])


def add_flow(doc):
    doc.add_heading("3. Flujo Operativo Actual", level=1)
    doc.add_heading("3.1 Venta Sin Proceso", level=2)
    for step in [
        "Vendedor crea cotizacion con cliente, perfil/tipo, cantidad, precio, entrega y condiciones.",
        "Contabilidad valida la cotizacion aceptada y la convierte en venta.",
        "La venta queda pendiente para bodega, sin descuento automatico de inventario.",
        "Bodega asigna lotes disponibles y cantidades.",
        "Bodega marca la venta como alistada; en ese momento se descuenta inventario.",
        "Bodega o contabilidad marca la venta como despachada.",
        "Contabilidad registra pagos o abonos hasta cerrar el saldo.",
    ]:
        add_numbered(doc, step)

    doc.add_heading("3.2 Venta Con Proceso Y Mezcla", level=2)
    for step in [
        "Vendedor cotiza el perfil solicitado, sin seleccionar lotes.",
        "Contabilidad convierte la cotizacion aceptada en venta.",
        "Bodega revisa la venta y solicita proceso si el cafe debe prepararse.",
        "Laboratorio confirma que el proceso inicio, registra ubicacion y fecha estimada de regreso.",
        "Al confirmar inicio se descuentan los lotes que entran a proceso.",
        "Cuando el proceso fisico termina, laboratorio lo marca como pendiente de examen.",
        "Laboratorio registra los examenes finales y crea el lote PROC.",
        "Laboratorio define si requiere mezcla final y registra porcentajes por lote.",
        "Bodega ve la orden de mezcla, la imprime si hace falta y alista el pedido.",
        "Al alistar o despachar se conserva el historial operativo y de pagos.",
    ]:
        add_numbered(doc, step)


def add_test_cases(doc):
    doc.add_heading("4. Guia Paso A Paso Para Verificar Funcionalidades", level=1)
    doc.add_paragraph(
        "Use datos de prueba o movimientos reales controlados. En cada caso se debe anotar si el resultado fue correcto, "
        "si hubo error o si el flujo se sintio lento/confuso para la persona que lo probo."
    )

    test_cases = [
        ("COM-01", "Vendedor / Comercial", "Crear cliente rapido", "Entrar a Comercial, abrir Crear cliente rapido, registrar nombre, telefono y direccion, guardar.", "El cliente queda creado, seleccionado en la cotizacion y disponible en la lista de clientes."),
        ("COM-02", "Vendedor / Comercial", "Crear cotizacion", "Seleccionar cliente, perfil o tipo, cantidad, precio, moneda, entrega y guardar.", "La cotizacion aparece en Historial con codigo, cliente, tipo, estado y total."),
        ("COM-03", "Vendedor / Comercial", "Filtrar cotizaciones", "Usar filtros Todas, Borradores, Enviadas, Aceptadas y Anuladas.", "La lista cambia segun el estado y los conteos coinciden con los registros."),
        ("CON-01", "Contabilidad / Comercial", "Aceptar cotizacion", "Seleccionar una cotizacion y cambiarla a Aceptada.", "La cotizacion queda aceptada y disponible para convertir a venta."),
        ("CON-02", "Contabilidad / Comercial", "Convertir a venta pagada", "Convertir cotizacion aceptada con estado Pagada y valor total pagado.", "Se crea venta pagada, pendiente para bodega, sin descontar inventario aun."),
        ("CON-03", "Contabilidad / Comercial", "Convertir con abono", "Convertir cotizacion aceptada con Pago parcial, valor abonado y fecha estimada de pago.", "La venta queda con saldo pendiente y estado de pago parcial."),
        ("VEN-01", "Ventas / Contabilidad", "Filtrar ventas por pago", "Entrar a Ventas y filtrar Pendientes, Parciales y Pagadas.", "La tabla muestra solo ventas del estado de pago seleccionado."),
        ("VEN-02", "Ventas / Contabilidad", "Registrar varios abonos", "Seleccionar una venta parcial, registrar un abono, luego registrar otro.", "Cada pago queda en historial, el saldo baja y al llegar a cero queda Pagada."),
        ("DOC-01", "Documentos", "Imprimir factura con pagos", "Abrir documento de venta con pagos y generar vista imprimible.", "La factura muestra historial de pagos con fecha, metodo, referencia, valor, notas, total pagado y saldo."),
        ("BOD-01", "Bodega", "Ver pendientes por prioridad", "Entrar a Pendientes y revisar ordenes por hacer.", "Se ven ventas activas ordenadas por urgencia, prioridad y fecha de entrega."),
        ("BOD-02", "Bodega", "Cambiar prioridad", "Seleccionar una venta pendiente y cambiar prioridad alta/media/baja.", "La prioridad se actualiza y la venta se reordena si corresponde."),
        ("BOD-03", "Bodega", "Asignar lote disponible", "Seleccionar venta, elegir producto vendido, lote disponible y cantidad, guardar asignacion.", "El lote queda asociado, pero inventario no se descuenta hasta marcar Alistada."),
        ("BOD-04", "Bodega", "Alistar venta", "Con lotes asignados, marcar la venta como Alistada.", "El inventario se descuenta y la venta queda lista para despacho."),
        ("BOD-05", "Bodega", "Despachar venta", "Seleccionar una venta alistada y marcar Despachada.", "La venta queda despachada y no se descuenta inventario nuevamente."),
        ("PRO-01", "Bodega / Procesos", "Crear solicitud de proceso", "Crear proceso asociado a una venta con lotes y cantidades.", "El proceso queda solicitado y no descuenta inventario hasta que laboratorio confirme inicio."),
        ("LAB-01", "Laboratorio / Procesos", "Confirmar inicio de proceso", "Entrar a Laboratorio, pestaña Procesos, seleccionar proceso pendiente, registrar ubicacion y fecha estimada, confirmar inicio.", "El proceso queda En proceso y se descuentan los lotes usados."),
        ("LAB-02", "Laboratorio / Procesos", "Marcar proceso terminado fisicamente", "Seleccionar proceso En proceso y marcar Pendiente de examen.", "El proceso queda pendiente de laboratorio; no aparece como lote disponible aun."),
        ("LAB-03", "Laboratorio / Procesos", "Finalizar examen y crear PROC", "Registrar perfil, peso final, humedad, catacion, score y guardar.", "Se crea lote PROC disponible y la venta asociada puede avanzar a ensamble/alistamiento."),
        ("LAB-04", "Laboratorio / Mezclas", "Crear orden de mezcla", "Seleccionar venta lista para ensamble, elegir producto, categoria, lote y porcentaje; guardar.", "La mezcla queda guardada y bodega puede ver/imprimir la orden."),
        ("BOD-06", "Bodega", "Imprimir orden de mezcla", "Abrir venta con mezcla definida e imprimir orden.", "El documento muestra productos, lotes, porcentajes y kg estimados."),
        ("MUE-01", "Muestras", "Crear solicitud de muestra", "Como vendedor/admin/contabilidad, crear muestra con solicitante, telefono, cafe, cantidad, precio opcional y fechas.", "La solicitud aparece registrada con estado Solicitada."),
        ("MUE-02", "Muestras", "Filtrar muestras", "Usar filtros Solicitadas, En preparacion, Listas, Entregadas y Canceladas.", "La lista muestra solo el estado seleccionado y deja primero lo pendiente."),
        ("MUE-03", "Muestras", "Cambiar estado de muestra", "Como usuario muestras, cambiar Solicitada a En preparacion, Lista y Entregada.", "El estado cambia y queda la ultima gestion registrada."),
        ("DASH-01", "Dashboard", "Revisar dashboard por rol", "Entrar con bodega, laboratorio, contabilidad y vendedor.", "Las tarjetas principales aparecen ordenadas segun el trabajo de cada rol."),
    ]
    add_table(
        doc,
        ["Codigo", "Modulo / Rol", "Caso", "Pasos", "Resultado esperado"],
        test_cases,
        [Inches(0.65), Inches(1.25), Inches(1.25), Inches(2.15), Inches(1.95)],
    )


def add_review_template(doc):
    doc.add_heading("5. Formato Para Registrar Hallazgos", level=1)
    doc.add_paragraph(
        "Durante las pruebas, se recomienda llenar esta tabla con cada problema, duda o mejora detectada."
    )
    rows = [("", "", "", "", "") for _ in range(8)]
    add_table(
        doc,
        ["Fecha", "Usuario / rol", "Modulo", "Hallazgo o mejora", "Prioridad"],
        rows,
        [Inches(0.9), Inches(1.2), Inches(1.1), Inches(3.0), Inches(0.9)],
    )


def add_closing(doc):
    doc.add_heading("6. Criterios De Aprobacion Inicial", level=1)
    for item in [
        "El vendedor puede crear clientes, cotizaciones y muestras sin depender de otros modulos.",
        "Contabilidad puede convertir ventas, registrar pagos y validar documentos con historial de abonos.",
        "Bodega puede ver pedidos pendientes, priorizar, asignar lotes, alistar, despachar e imprimir ordenes.",
        "Laboratorio puede controlar procesos, bloquear regreso a bodega sin examen y definir mezclas.",
        "El inventario no presenta descuentos dobles ni movimientos negativos inesperados.",
        "Cada rol entiende que debe hacer al entrar a su pantalla principal.",
    ]:
        add_bullet(doc, item)


def build_document():
    doc = Document()
    setup_styles(doc)
    add_title(doc)
    add_summary(doc)
    add_changes(doc)
    add_flow(doc)
    add_test_cases(doc)
    add_review_template(doc)
    add_closing(doc)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
