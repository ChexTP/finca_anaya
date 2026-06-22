from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Guia_pruebas_bodega_muestras_laboratorio.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Body")
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [Inches(6.5)])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}: ")
    set_run(r, bold=True, color="1F4D78")
    p.add_run(text)
    doc.add_paragraph()


def add_test_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_width(table, [Inches(0.9), Inches(2.15), Inches(1.85), Inches(1.6)])
    headers = ["Prueba", "Accion a realizar", "Resultado esperado", "Observaciones"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "E8EEF5")
        run = cell.paragraphs[0].add_run(header)
        set_run(run, bold=True, color="0B2545")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.style = doc.styles["Table Body"]

    doc.add_paragraph()
    return table


def add_feedback_table(doc):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_table_width(table, [Inches(1.0), Inches(1.35), Inches(1.55), Inches(1.55), Inches(1.05)])
    headers = ["Modulo", "Pantalla", "Que se probo", "Resultado/feedback", "Prioridad"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "E8EEF5")
        run = cell.paragraphs[0].add_run(header)
        set_run(run, bold=True, color="0B2545")

    for _ in range(8):
        row = table.add_row().cells
        for cell in row:
            cell.text = ""
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles["Table Body"]


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    body = styles.add_style("Body", 1)
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.25

    table_body = styles.add_style("Table Body", 1)
    table_body.font.name = "Calibri"
    table_body.font.size = Pt(9)
    table_body.paragraph_format.space_after = Pt(0)
    table_body.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_document():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Guia de pruebas funcionales")
    set_run(run, size=22, color="0B2545", bold=True)
    subtitle = doc.add_paragraph(style="Body")
    subtitle.add_run("Modulos: Bodega, Muestras y Laboratorio | Sistema Finca Anaya")
    meta = doc.add_paragraph(style="Body")
    meta.add_run("Objetivo: validar con datos reales del trabajo diario si el sistema responde al flujo operativo esperado.")

    add_note(
        doc,
        "Uso recomendado",
        "Realizar las pruebas con datos de movimientos reales o simulados de la empresa, anotar cualquier diferencia contra la forma actual de trabajo y reportar capturas cuando algo no sea claro.",
    )

    add_heading(doc, "1. Acceso y preparacion", 1)
    add_body(doc, "Antes de probar, confirmar que el sistema este abierto en el navegador y que se ingrese con el usuario correcto segun el modulo.")
    add_test_table(
        doc,
        [
            ("A1", "Ingresar con usuario bodega / bodega123.", "El sistema abre el modulo Bodega.", ""),
            ("A2", "Ingresar con usuario laboratorio / laboratorio123.", "El sistema abre el modulo Laboratorio.", ""),
            ("A3", "Ingresar con usuario muestras / muestras123.", "El sistema abre el modulo Muestras.", ""),
            ("A4", "Ingresar con vendedor1 / vendedor123.", "El vendedor puede crear solicitudes de muestras y cotizaciones.", ""),
        ],
    )

    add_heading(doc, "2. Modulo Bodega", 1)
    add_body(doc, "Este modulo valida la recepcion inicial del cafe. Aqui se registra proveedor, peso, empaque, humedad, merma por trilla, examen visual y observaciones.")
    add_heading(doc, "Datos que debe preparar la persona que prueba", 2)
    add_bullets(
        doc,
        [
            "Nombre, telefono, direccion y zona de procedencia del proveedor o finca.",
            "Tipo de cafe recibido: pergamino, trillado, procesado o especial.",
            "Peso bruto en kg, tipo de empaque, cantidad de empaques y si tiene bolsa interna.",
            "Humedad %, merma por trilla %, resultado del examen visual y defecto visual % si aplica.",
            "Observaciones visuales y comentario inicial del lote.",
        ],
    )
    add_heading(doc, "Pruebas sugeridas de Bodega", 2)
    add_test_table(
        doc,
        [
            (
                "B1",
                "Crear un proveedor nuevo desde Proveedor rapido.",
                "El proveedor queda guardado y aparece disponible para seleccionar al recibir cafe.",
                "",
            ),
            (
                "B2",
                "Ingresar cafe aprobado visualmente con peso bruto, empaque, humedad y merma por trilla.",
                "El sistema calcula el peso neto, genera un codigo LOT y lo muestra como pendiente de laboratorio.",
                "",
            ),
            (
                "B3",
                "Marcar Tiene bolsa interna y revisar el peso estimado.",
                "El peso neto descuenta el peso del empaque y 0.05 kg por cada bolsa interna.",
                "",
            ),
            (
                "B4",
                "Ingresar Merma por trilla % con un valor entre 0 y 100.",
                "El lote se guarda y la merma aparece en la tabla Pendientes de laboratorio.",
                "",
            ),
            (
                "B5",
                "Intentar ingresar Merma por trilla % mayor a 100.",
                "El sistema debe rechazar el registro y mostrar un mensaje de validacion.",
                "",
            ),
            (
                "B6",
                "Ingresar cafe con examen visual rechazado.",
                "El sistema no debe enviarlo a pendientes de laboratorio; debe registrar que no continua el flujo operativo.",
                "",
            ),
            (
                "B7",
                "Intentar guardar sin proveedor, tipo de cafe, peso bruto, empaque o examen visual.",
                "El sistema debe pedir los datos obligatorios y no crear el lote incompleto.",
                "",
            ),
        ],
    )
    add_note(
        doc,
        "Criterio de aceptacion en Bodega",
        "El responsable debe poder registrar un cafe recibido sin usar hojas externas para los datos principales de trazabilidad: proveedor, origen, humedad, merma, peso y examen visual.",
    )

    add_heading(doc, "3. Modulo Laboratorio", 1)
    add_body(doc, "Este modulo valida los lotes recibidos por bodega y los procesos pendientes. La pantalla inicial debe mostrar primero los lotes pendientes de revision.")
    add_heading(doc, "Pruebas sugeridas de Laboratorio", 2)
    add_test_table(
        doc,
        [
            (
                "L1",
                "Entrar al modulo Laboratorio despues de crear un lote aprobado en Bodega.",
                "El lote aparece en la lista principal de lotes pendientes.",
                "",
            ),
            (
                "L2",
                "Seleccionar un lote pendiente.",
                "El formulario indica el codigo seleccionado y carga la humedad registrada como referencia.",
                "",
            ),
            (
                "L3",
                "Verificar que se muestre la merma por trilla del lote.",
                "La persona de laboratorio puede ver el dato registrado en Bodega antes de aprobar.",
                "",
            ),
            (
                "L4",
                "Aprobar un lote con humedad y datos de catacion completos: aroma, fragancia, sabor, acidez, dulzor, cuerpo, balance, uniformidad, residual, taza limpia y score.",
                "El lote cambia a aprobado y queda pendiente para que contabilidad registre la compra/pago.",
                "",
            ),
            (
                "L5",
                "Intentar aprobar un lote sin completar score o campos de catacion.",
                "El sistema debe impedir la aprobacion y pedir los datos faltantes.",
                "",
            ),
            (
                "L6",
                "Rechazar un lote desde laboratorio.",
                "El lote queda rechazado y no pasa al inventario disponible.",
                "",
            ),
            (
                "L7",
                "Entrar a la pestana Procesos del menu interno de Laboratorio.",
                "Debe mostrar procesos pendientes para revisar y crear el lote procesado final cuando aplique.",
                "",
            ),
            (
                "L8",
                "Finalizar un proceso con peso final, perfil, humedad, score y notas.",
                "El sistema crea un lote PROC disponible/registrado con los datos de laboratorio del cafe procesado.",
                "",
            ),
        ],
    )
    add_note(
        doc,
        "Criterio de aceptacion en Laboratorio",
        "El laboratorio debe poder decidir si un lote entra o no a la empresa y debe dejar trazables los datos de calidad que antes podian perderse en notas fisicas.",
    )

    add_heading(doc, "4. Modulo Muestras", 1)
    add_body(doc, "Este modulo permite crear solicitudes de muestras para clientes y que la persona encargada de muestras gestione su estado.")
    add_heading(doc, "Roles que intervienen", 2)
    add_bullets(
        doc,
        [
            "Vendedor: crea solicitudes y ve sus propias solicitudes.",
            "Administrador y contabilidad: crean y ven solicitudes.",
            "Usuario muestras: recibe las solicitudes y actualiza estados.",
        ],
    )
    add_heading(doc, "Pruebas sugeridas de Muestras", 2)
    add_test_table(
        doc,
        [
            (
                "M1",
                "Ingresar como vendedor y crear una muestra gratis dejando el precio vacio.",
                "La solicitud se crea con precio Gratis y queda en estado Solicitada.",
                "",
            ),
            (
                "M2",
                "Crear una muestra con precio en COP o USD.",
                "La solicitud muestra la moneda y el valor indicado.",
                "",
            ),
            (
                "M3",
                "Crear una muestra seleccionando un perfil de cafe.",
                "La solicitud queda asociada al perfil seleccionado.",
                "",
            ),
            (
                "M4",
                "Crear una muestra seleccionando un tipo de cafe o escribiendo descripcion libre.",
                "La solicitud guarda la referencia del cafe solicitado aun si no hay perfil definido.",
                "",
            ),
            (
                "M5",
                "Ingresar como usuario muestras.",
                "La pantalla muestra las solicitudes creadas por vendedores, admin o contabilidad.",
                "",
            ),
            (
                "M6",
                "Cambiar el estado de una muestra a En preparacion, Lista, Entregada o Cancelada.",
                "El sistema pide confirmacion y actualiza el estado en la lista.",
                "",
            ),
            (
                "M7",
                "Agregar una nota al cambiar estado.",
                "La nota queda visible dentro de la solicitud para seguimiento interno.",
                "",
            ),
            (
                "M8",
                "Intentar crear una muestra sin nombre, telefono, cantidad o cafe solicitado.",
                "El sistema debe impedir el registro y mostrar validacion.",
                "",
            ),
        ],
    )
    add_note(
        doc,
        "Criterio de aceptacion en Muestras",
        "La persona encargada debe poder saber que muestras fueron solicitadas, para quien, que cafe pidieron, cuanta cantidad, si se cobra o no, y en que estado va la preparacion.",
    )

    add_heading(doc, "5. Flujo de prueba recomendado", 1)
    add_numbered(
        doc,
        [
            "Crear o seleccionar un proveedor real y registrar un cafe recibido desde Bodega.",
            "Confirmar que el lote aparezca pendiente en Bodega y luego en Laboratorio.",
            "Entrar a Laboratorio, revisar el lote y aprobarlo o rechazarlo con datos reales.",
            "Crear una solicitud de muestra como vendedor con un perfil o tipo de cafe que la empresa maneje.",
            "Entrar como usuario muestras y cambiar el estado de esa solicitud hasta Lista o Entregada.",
            "Anotar si algun dato que usan en el trabajo diario no tiene campo claro en el sistema.",
        ],
    )

    add_heading(doc, "6. Registro de feedback", 1)
    add_body(doc, "Usar esta tabla para devolver observaciones concretas. Idealmente cada comentario debe indicar modulo, pantalla, dato probado y resultado esperado.")
    add_feedback_table(doc)

    add_heading(doc, "7. Preguntas clave para el evaluador", 1)
    add_bullets(
        doc,
        [
            "Que campo hizo falta para registrar un caso real?",
            "Hubo algun paso que parecio lento o confuso?",
            "El nombre de los estados coincide con la forma en que hablan dentro de la empresa?",
            "Los datos visibles en cada pantalla son suficientes para trabajar sin preguntar por fuera?",
            "Hay algun permiso que permita ver informacion que ese rol no deberia ver?",
            "Que cambio tendria mayor impacto para usar el sistema en el dia a dia?",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Guia de pruebas - Sistema Finca Anaya")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
