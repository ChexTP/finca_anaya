from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/guia-cambios-pruebas-flujo-actual-2026-07-02.docx"
BLUE = "2E74B5"
DARK_BLUE = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "EAF6EE"
GREEN = "256B45"


def set_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill, border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        borders.append(left)
        p_pr.append(borders)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "decimal")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "271")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    numbering.append(number)
    return num_id


def add_numbered_steps(doc, steps):
    num_id = create_numbering(doc)
    for step in steps:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        number = OxmlElement("w:numId")
        number.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(number)
        paragraph._p.get_or_add_pPr().append(num_pr)
        set_font(paragraph.add_run(step))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Test Result" not in doc.styles:
        result_style = doc.styles.add_style("Test Result", WD_STYLE_TYPE.PARAGRAPH)
    else:
        result_style = doc.styles["Test Result"]
    result_style.font.name = "Calibri"
    result_style.font.size = Pt(10.5)
    result_style.paragraph_format.left_indent = Inches(0.12)
    result_style.paragraph_format.right_indent = Inches(0.08)
    result_style.paragraph_format.space_before = Pt(4)
    result_style.paragraph_format.space_after = Pt(8)
    result_style.paragraph_format.line_spacing = 1.2

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("FINCA ANAYA  |  GUIA OPERATIVA DE VALIDACION")
    set_font(run, size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_title_block(doc):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    set_font(kicker.add_run("GUIA DE IMPLEMENTACION Y PRUEBAS"), size=10, bold=True, color=GREEN)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("Cambios y flujo operativo actual"), size=25, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(
        subtitle.add_run("Inventario, laboratorio, ventas, procesos, muestras y control de necesidades"),
        size=13,
        color=MUTED,
    )

    for label, value in [
        ("Sistema", "Finca Anaya"),
        ("Version del documento", "2 de julio de 2026"),
        ("Objetivo", "Validar el flujo acordado en la reunion del 30 de junio de 2026"),
    ]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_font(paragraph.add_run(f"{label}: "), bold=True, color=DARK_BLUE)
        set_font(paragraph.add_run(value), color="344054")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(12)
    intro.paragraph_format.space_after = Pt(12)
    shade_paragraph(intro, LIGHT_BLUE, BLUE)
    set_font(intro.add_run("Alcance. "), bold=True, color=DARK_BLUE)
    set_font(
        intro.add_run(
            "Este documento resume los cambios implementados y define las pruebas que debe ejecutar la empresa. "
            "Cada caso indica los pasos y el resultado que debe observarse antes de aprobar el flujo."
        ),
        color="344054",
    )


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_font(paragraph.add_run(item))


def add_changes(doc):
    doc.add_heading("1. Cambios implementados", level=1)
    sections = [
        ("1.1 Bodega e inventario", [
            "Los metodos de proceso disponibles son Lavado, Natural y Semilavado.",
            "La recepcion registra categoria Regional, Varietal o Exotico y una clasificacion o variedad en texto libre.",
            "La fecha exacta de llegada a bodega, la humedad y el factor de rendimiento son obligatorios.",
            "Bodega ya no aprueba ni rechaza el cafe al recibirlo: todo lote queda pendiente de laboratorio.",
            "El inventario muestra categoria, clasificacion y fecha de llegada.",
        ]),
        ("1.2 Ventas y cotizaciones", [
            "Una cotizacion puede contener varios cafes con caracteristicas y precios diferentes.",
            "Cada item registra Excelso o Pergamino, metodo de proceso, perfil o descripcion, variedad, cantidad y precio.",
            "La fecha estimada de entrega es obligatoria.",
            "Los nuevos datos aparecen en el detalle comercial, documentos y ordenes operativas.",
        ]),
        ("1.3 Procesos y laboratorio", [
            "Bodega y administracion pueden crear solicitudes de proceso.",
            "Solo administracion confirma el inicio, registra ubicacion y fecha de regreso, y envia el proceso terminado a examen.",
            "El inventario de entrada se descuenta cuando administracion confirma el inicio.",
            "Laboratorio solo recibe lotes y procesos pendientes de examen y registra los resultados tecnicos.",
            "Al finalizar el examen de un proceso se crea el lote PROC correspondiente.",
        ]),
        ("1.4 Muestras", [
            "Las cantidades se registran y muestran en gramos.",
            "La cantidad esta separada visualmente del precio y la moneda.",
            "El precio sigue siendo opcional; sin precio, la muestra aparece como Gratis.",
        ]),
        ("1.5 Dashboard y documentacion", [
            "Administracion, contabilidad y bodega ven un tablero de necesidades de cafe.",
            "El tablero compara cantidad solicitada, inventario compatible y cantidad por conseguir.",
            "No realiza conversiones ni calcula rendimientos de produccion.",
            "La guia del backend fue actualizada con campos, permisos y respuestas vigentes.",
        ]),
    ]
    for heading, bullets in sections:
        doc.add_heading(heading, level=2)
        add_bullets(doc, bullets)


def add_flow(doc):
    doc.add_heading("2. Flujo operativo vigente", level=1)
    flows = [
        ("2.1 Ingreso de cafe", [
            "Bodega registra proveedor, proceso, fecha de llegada, peso, embalaje, humedad, factor, categoria y variedad.",
            "El sistema calcula el peso neto y crea un codigo LOT.",
            "El lote queda pendiente de laboratorio sin aprobacion previa de bodega.",
            "Laboratorio registra la evaluacion sensorial y decide aprobar o rechazar.",
            "Si se aprueba, contabilidad registra la compra y el lote pasa a inventario disponible.",
        ]),
        ("2.2 Venta con inventario disponible", [
            "El vendedor crea una cotizacion con uno o varios cafes y fecha de entrega.",
            "Contabilidad o administracion acepta y convierte la cotizacion en venta.",
            "Bodega prioriza la venta y asigna lotes a cada producto.",
            "El inventario se descuenta cuando la venta se marca como alistada.",
            "Bodega marca la salida como despachada y contabilidad mantiene el control de pagos.",
        ]),
        ("2.3 Venta que requiere proceso", [
            "Bodega solicita el proceso asociado a la venta y selecciona lotes y cantidades.",
            "Administracion confirma el inicio; en ese momento se descuenta el cafe utilizado.",
            "Administracion registra cuando el proceso fisico termina y lo envia a examen.",
            "Laboratorio registra peso final, humedad, catacion, score y perfil.",
            "El sistema crea el lote PROC y la venta puede continuar con el alistamiento.",
        ]),
    ]
    for heading, steps in flows:
        doc.add_heading(heading, level=2)
        add_numbered_steps(doc, steps)


def add_test(doc, code, title, role, steps, expected):
    doc.add_heading(f"{code} - {title}", level=3)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(5)
    set_font(meta.add_run("Rol: "), bold=True, color=DARK_BLUE)
    set_font(meta.add_run(role), italic=True, color=MUTED)
    add_numbered_steps(doc, steps)
    result = doc.add_paragraph(style="Test Result")
    shade_paragraph(result, LIGHT_GREEN, GREEN)
    set_font(result.add_run("Resultado esperado: "), bold=True, color=GREEN)
    set_font(result.add_run(expected), color="344054")


def add_tests(doc):
    doc.add_heading("3. Guia de pruebas funcionales", level=1)
    paragraph = doc.add_paragraph()
    set_font(
        paragraph.add_run(
            "Ejecute los casos en el orden indicado cuando necesite validar el flujo completo. Use datos de prueba identificables y registre cualquier diferencia observada."
        )
    )

    tests = [
        ("BOD-01", "Recibir cafe", "Bodega", [
            "Entrar al modulo Bodega y crear o seleccionar un proveedor.",
            "Registrar metodo de proceso, fecha de llegada, peso, embalaje, humedad y factor de rendimiento.",
            "Registrar categoria, clasificacion o variedad y observaciones.",
            "Presionar Registrar cafe.",
        ], "Se genera un codigo LOT y el cafe queda Pendiente de laboratorio. No aparece una decision visual. El formulario no permite guardar sin fecha, humedad o factor."),
        ("LAB-01", "Evaluar lote recibido", "Laboratorio", [
            "Entrar a Laboratorio y seleccionar el lote pendiente.",
            "Registrar humedad, atributos de catacion, score y notas.",
            "Seleccionar Aprobar o Rechazar y confirmar.",
        ], "Si se aprueba, el lote queda disponible inmediatamente aunque el proveedor aun no haya cobrado. Si se rechaza, queda como historico y no entra al inventario disponible."),
        ("CON-01", "Registrar pago de lote aprobado", "Contabilidad", [
            "Entrar a Inventario y seleccionar un lote aprobado pendiente de pago.",
            "Registrar precio por kg, metodo, referencia y fecha de pago.",
            "Confirmar el registro del pago.",
        ], "El pago y el costo quedan registrados sin modificar los kilos disponibles del lote."),
        ("COM-01", "Crear cotizacion con varios cafes", "Vendedor", [
            "Entrar a Comercial, seleccionar cliente, moneda y fecha de entrega.",
            "Registrar el primer cafe con presentacion, proceso, perfil o descripcion, variedad, cantidad y precio.",
            "Presionar Agregar otro cafe y registrar un segundo item diferente.",
            "Guardar la cotizacion.",
        ], "Se crea una sola cotizacion con ambos items, subtotales correctos y costo de envio incluido. No se descuenta inventario. Sin fecha de entrega no debe permitir guardar."),
        ("CON-02", "Convertir cotizacion en venta", "Administracion o Contabilidad", [
            "Abrir la cotizacion creada y cambiarla a Aceptada.",
            "Convertirla en venta y registrar el estado de pago correspondiente.",
            "Consultar la venta generada.",
        ], "Se genera un codigo VEN, se conservan todos los cafes y la orden queda pendiente para bodega sin descuento inmediato de inventario."),
        ("BOD-02", "Asignar lotes y despachar venta", "Bodega", [
            "Entrar a Pendientes, seleccionar la venta y definir su prioridad.",
            "Asignar un lote y cantidad a cada producto.",
            "Marcar la venta como Alistada y comprobar el inventario.",
            "Marcarla posteriormente como Despachada.",
        ], "Los lotes quedan asociados, el descuento ocurre una sola vez al alistar y la venta deja de aparecer como pendiente cuando se despacha."),
        ("PRO-01", "Solicitar proceso", "Bodega", [
            "Entrar a Procesos y seleccionar la venta correspondiente.",
            "Seleccionar los lotes y cantidades requeridas.",
            "Presionar Solicitar proceso.",
        ], "El proceso queda Pendiente y todavía no descuenta inventario. El usuario de bodega no puede iniciarlo."),
        ("PRO-02", "Iniciar proceso", "Administrador", [
            "Entrar a Procesos y localizar la solicitud pendiente.",
            "Presionar Iniciar proceso.",
            "Registrar ubicacion y fecha estimada de regreso y confirmar.",
        ], "El proceso pasa a En proceso y se descuentan exactamente las cantidades seleccionadas de los lotes origen."),
        ("PRO-03", "Enviar proceso a examen", "Administrador", [
            "Seleccionar el proceso En proceso.",
            "Presionar Enviar a examen y confirmar.",
        ], "El proceso pasa a Pendiente de laboratorio y aparece para el usuario de laboratorio. Bodega no puede ejecutar esta accion."),
        ("LAB-02", "Finalizar examen de proceso", "Laboratorio", [
            "Entrar a Laboratorio y abrir el proceso pendiente de examen.",
            "Registrar perfil, peso final, humedad, catacion completa y score.",
            "Confirmar la creacion del lote procesado.",
        ], "El proceso queda Finalizado, se crea un codigo PROC y el nuevo lote queda disponible para continuar la venta."),
        ("MUE-01", "Registrar muestras en gramos", "Vendedor, Administracion o Contabilidad", [
            "Crear una solicitud de 250 gramos sin precio.",
            "Crear otra solicitud con cantidad en gramos y precio.",
            "Consultar ambas solicitudes.",
        ], "Las cantidades aparecen en gramos. La primera muestra aparece como Gratis y la segunda muestra moneda y precio sin confundirlos con la cantidad."),
        ("DASH-01", "Revisar necesidades de cafe", "Administracion, Contabilidad o Bodega", [
            "Crear o conservar una venta activa con fecha de entrega.",
            "Entrar al Dashboard y ubicar Necesidades de cafe.",
            "Comparar solicitado, disponible y por conseguir.",
        ], "El tablero agrupa la demanda por presentacion, proceso, perfil y variedad, muestra la entrega mas proxima y no modifica el inventario ni calcula rendimientos."),
        ("DOC-01", "Validar documentos", "Vendedor, Administracion o Contabilidad", [
            "Abrir la cotizacion y la venta desde Documentos.",
            "Generar la vista imprimible o guardar como PDF.",
            "Revisar productos, totales y pagos registrados.",
        ], "Los documentos muestran todos los cafes, sus caracteristicas, cantidades, precios, totales y el historial de pagos cuando corresponde."),
    ]
    for test in tests:
        add_test(doc, *test)


def add_acceptance(doc):
    doc.add_heading("4. Criterios de aprobacion", level=1)
    items = [
        "Bodega registra lotes completos sin tomar decisiones de calidad que corresponden a laboratorio.",
        "Laboratorio solo recibe lotes y procesos que requieren examen.",
        "Bodega solicita procesos, pero solo administracion puede iniciarlos o enviarlos a examen.",
        "Una venta conserva correctamente todos sus items desde la cotizacion hasta el documento final.",
        "No existen descuentos dobles ni cantidades negativas inesperadas.",
        "El tablero de necesidades informa faltantes sin ejecutar movimientos de inventario.",
        "Las muestras se entienden claramente en gramos y el precio opcional no se confunde con la cantidad.",
    ]
    add_bullets(doc, items)

    doc.add_heading("5. Registro de hallazgos", level=1)
    for label in ["Fecha y usuario", "Modulo y caso", "Resultado obtenido", "Diferencia encontrada", "Prioridad sugerida"]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)
        set_font(paragraph.add_run(f"{label}: "), bold=True, color=DARK_BLUE)
        set_font(paragraph.add_run("____________________________________________________________"), color="98A2B3")


def build():
    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    add_changes(doc)
    add_flow(doc)
    add_tests(doc)
    add_acceptance(doc)
    doc.core_properties.title = "Cambios y pruebas del flujo operativo actual - Finca Anaya"
    doc.core_properties.subject = "Guia funcional para validacion del sistema"
    doc.core_properties.author = "ChexDev"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
