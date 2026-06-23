from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Guia_casos_uso_usuarios_sistema_anaya.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=None, bold=False, color=None):
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    body = doc.styles.add_style("Body Text Compact", 1)
    body.font.name = "Calibri"
    body.font.size = Pt(10.5)
    body.paragraph_format.space_after = Pt(5)
    body.paragraph_format.line_spacing = 1.18

    table_body = doc.styles.add_style("Table Body Compact", 1)
    table_body.font.name = "Calibri"
    table_body.font.size = Pt(8.7)
    table_body.paragraph_format.space_after = Pt(0)
    table_body.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.2

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15


def add_heading(doc, text, level=1):
    doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text):
    doc.add_paragraph(text, style="Body Text Compact")


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [Inches(6.5)])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}: ")
    set_run(r, bold=True, color="1F4D78")
    p.add_run(text)
    doc.add_paragraph()


def add_credentials_table(doc):
    rows = [
        ("Administrador", "admin", "admin123", "Acceso total al sistema."),
        ("Contabilidad", "contabilidad", "contabilidad123", "Ventas, pagos, compras, cuentas, reportes y documentos."),
        ("Bodega", "bodega", "bodega123", "Recepcion de cafe, procesos, inventario operativo y despachos."),
        ("Laboratorio", "laboratorio", "laboratorio123", "Revision de lotes recibidos y procesos finalizados."),
        ("Vendedor", "vendedor1", "vendedor123", "Clientes, cotizaciones, preventas, muestras y documentos comerciales."),
        ("Muestras", "muestras", "muestras123", "Gestion de solicitudes de muestras."),
    ]
    add_matrix(doc, ["Rol", "Usuario", "Clave", "Uso principal"], rows, [1.35, 1.25, 1.25, 2.65])


def add_matrix(doc, headers, rows, widths_in):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, [Inches(width) for width in widths_in])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "E8EEF5")
        run = cell.paragraphs[0].add_run(header)
        set_run(run, bold=True, color="0B2545")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.style = doc.styles["Table Body Compact"]
    doc.add_paragraph()


def add_use_cases(doc, role_title, purpose, cases):
    add_heading(doc, role_title, 1)
    add_body(doc, purpose)
    rows = []
    for code, module, action, expected, feedback in cases:
        rows.append((code, module, action, expected, feedback))
    add_matrix(
        doc,
        ["Caso", "Modulo", "Accion a probar", "Resultado esperado", "Observaciones"],
        rows,
        [0.65, 1.05, 2.25, 1.85, 0.7],
    )


def build_document():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Guia de casos de uso por usuario")
    set_run(run, size=22, color="0B2545", bold=True)
    add_body(doc, "Sistema web Finca Anaya - Guia para pruebas funcionales por rol.")
    add_body(doc, "Objetivo: que la persona encargada de probar pueda recorrer las funciones principales del sistema, verificar resultados esperados y registrar feedback concreto para ajustes.")

    add_note(
        doc,
        "Recomendacion",
        "Probar con datos cercanos a la operacion real de la empresa. Si un flujo no coincide con la forma diaria de trabajo, registrar el caso con detalle en la columna de observaciones.",
    )

    add_heading(doc, "Accesos de prueba", 1)
    add_credentials_table(doc)

    add_heading(doc, "Criterios generales de validacion", 1)
    add_bullets(
        doc,
        [
            "Cada rol debe ver solo los modulos que necesita para su trabajo.",
            "El sistema debe mostrar mensajes claros cuando falten datos obligatorios.",
            "Las cantidades de cafe no deben quedar negativas.",
            "Los cambios importantes deben pedir confirmacion cuando aplique.",
            "La informacion creada debe aparecer inmediatamente en las listas correspondientes.",
            "Los reportes y documentos deben reflejar los datos ingresados en ventas, pagos, inventario y cuentas.",
        ]
    )

    add_use_cases(
        doc,
        "Administrador",
        "El administrador valida el funcionamiento general, usuarios, reportes, documentos y acceso a todos los modulos.",
        [
            ("ADM-01", "Dashboard", "Ingresar como admin y revisar tarjetas/alertas.", "Debe ver resumen de inventario, ventas, cartera y alertas.", ""),
            ("ADM-02", "Usuarios", "Crear un vendedor nuevo.", "El vendedor queda disponible para iniciar sesion.", ""),
            ("ADM-03", "Usuarios", "Cambiar clave de un usuario.", "El usuario puede iniciar con la nueva clave.", ""),
            ("ADM-04", "Usuarios", "Desactivar y activar un usuario.", "El acceso queda bloqueado o permitido segun estado.", ""),
            ("ADM-05", "Clientes", "Crear y editar un cliente.", "Los cambios se guardan y aparecen en cotizaciones/ventas.", ""),
            ("ADM-06", "Inventario", "Revisar lotes disponibles y aprobados.", "Debe ver pesos, humedad, factor de rendimiento, clasificacion comercial, perfil/tipo y estado.", ""),
            ("ADM-07", "Comercial", "Crear una cotizacion/preventa.", "La cotizacion queda guardada en historial.", ""),
            ("ADM-08", "Comercial", "Convertir cotizacion aceptada en venta.", "Se crea venta y queda pendiente de alistamiento.", ""),
            ("ADM-09", "Ventas", "Revisar venta creada y pagos.", "Debe ver total, abonado, saldo y estado operativo.", ""),
            ("ADM-10", "Reportes", "Filtrar ventas por fechas, cliente, estado o vendedor.", "El reporte debe mostrar datos coherentes y permitir exportar CSV.", ""),
            ("ADM-11", "Backups", "Descargar backup de un modulo.", "Se descarga CSV y queda registro en historial.", ""),
            ("ADM-12", "Documentos", "Abrir documento de cotizacion o venta.", "Debe verse vista imprimible/descargable.", ""),
        ],
    )

    add_use_cases(
        doc,
        "Contabilidad",
        "Contabilidad valida compras, pagos, ventas, cuentas por pagar, documentos y reportes financieros simples.",
        [
            ("CON-01", "Inventario", "Seleccionar lote aprobado por laboratorio.", "Debe permitir registrar precio por kg y pago del lote.", ""),
            ("CON-02", "Inventario", "Registrar compra/pago de lote.", "El lote pasa a disponible en inventario.", ""),
            ("CON-03", "Comercial", "Revisar cotizacion aceptada.", "Debe poder convertirla en venta.", ""),
            ("CON-04", "Comercial", "Convertir cotizacion en venta con pago total.", "La venta queda pagada y lista para alistamiento.", ""),
            ("CON-05", "Comercial", "Convertir cotizacion con abono parcial.", "La venta queda con saldo pendiente y fecha estimada de pago.", ""),
            ("CON-06", "Ventas", "Registrar abono a una venta pendiente.", "El saldo se recalcula y el pago queda en historial.", ""),
            ("CON-07", "Ventas", "Validar venta pagada.", "Cuando el saldo llega a cero debe quedar como pagada.", ""),
            ("CON-08", "Cuentas por pagar", "Crear cuenta por pagar manual.", "La cuenta queda pendiente con categoria y tercero/proveedor.", ""),
            ("CON-09", "Cuentas por pagar", "Registrar pago de cuenta por pagar.", "El saldo baja y el estado cambia si se paga completa.", ""),
            ("CON-10", "Reportes", "Revisar ventas, cartera, cuentas por pagar y utilidad estimada.", "Los valores deben coincidir con los movimientos registrados.", ""),
            ("CON-11", "Documentos", "Generar documento de venta/cotizacion.", "Debe poder imprimirse o guardarse como PDF desde el navegador.", ""),
            ("CON-12", "Backups", "Exportar ventas, pagos o cuentas por pagar.", "Debe descargar CSV de respaldo.", ""),
        ],
    )

    add_use_cases(
        doc,
        "Bodega",
        "Bodega valida recepcion de cafe, calculo de peso neto, trazabilidad inicial, seleccion para procesos y despacho de ventas.",
        [
            ("BOD-01", "Bodega", "Crear proveedor rapido.", "El proveedor queda disponible para seleccionar.", ""),
            ("BOD-02", "Bodega", "Ingresar cafe recibido aprobado visualmente.", "Se genera codigo LOT y queda pendiente de laboratorio.", ""),
            ("BOD-03", "Bodega", "Registrar peso bruto, empaque y bolsa interna.", "El peso neto debe descontar empaque y bolsa interna.", ""),
            ("BOD-04", "Bodega", "Registrar humedad, factor de rendimiento y clasificacion comercial.", "Los datos quedan visibles en pendientes de laboratorio.", ""),
            ("BOD-05", "Bodega", "Intentar guardar factor de rendimiento negativo.", "El sistema debe rechazar el dato.", ""),
            ("BOD-06", "Bodega", "Ingresar lote rechazado visualmente.", "No debe continuar a laboratorio como pendiente operativo.", ""),
            ("BOD-07", "Procesos", "Crear proceso seleccionando lotes disponibles y cantidades.", "El proceso queda pendiente/en proceso y descuenta disponibilidad operativa.", ""),
            ("BOD-08", "Procesos", "Intentar usar mas cafe del disponible.", "El sistema debe impedirlo.", ""),
            ("BOD-09", "Ventas", "Ver venta pendiente de alistamiento.", "Debe ver lotes/cantidades que debe preparar.", ""),
            ("BOD-10", "Ventas", "Cambiar venta a alistada.", "El estado cambia y queda lista para despacho.", ""),
            ("BOD-11", "Ventas", "Cambiar venta a despachada.", "El sistema registra que bodega ya despacho.", ""),
        ],
    )

    add_use_cases(
        doc,
        "Laboratorio",
        "Laboratorio valida calidad de lotes recibidos y registra datos finales de procesos de cafe especial.",
        [
            ("LAB-01", "Laboratorio", "Entrar y ver lotes pendientes.", "La pantalla principal muestra lotes enviados por bodega.", ""),
            ("LAB-02", "Laboratorio", "Seleccionar un lote pendiente.", "El formulario queda asociado al codigo seleccionado.", ""),
            ("LAB-03", "Laboratorio", "Ver humedad y factor de rendimiento registrado en bodega.", "Debe ver esos datos antes de decidir.", ""),
            ("LAB-04", "Laboratorio", "Aprobar lote con catacion completa y score.", "El lote pasa a aprobado para que contabilidad lo compre/pague.", ""),
            ("LAB-05", "Laboratorio", "Intentar aprobar sin score o catacion completa.", "El sistema debe impedir aprobacion.", ""),
            ("LAB-06", "Laboratorio", "Rechazar lote por laboratorio.", "El lote no debe pasar a inventario disponible.", ""),
            ("LAB-07", "Laboratorio/Procesos", "Abrir pestana Procesos.", "Debe ver procesos pendientes de finalizar.", ""),
            ("LAB-08", "Laboratorio/Procesos", "Finalizar proceso con perfil, peso final, humedad y score.", "Se crea lote PROC con datos exactos del cafe procesado.", ""),
            ("LAB-09", "Laboratorio/Procesos", "Intentar finalizar proceso sin datos obligatorios.", "El sistema debe mostrar validacion.", ""),
        ],
    )

    add_use_cases(
        doc,
        "Vendedor",
        "El vendedor valida clientes, cotizaciones/preventas, solicitudes de muestras y documentos comerciales.",
        [
            ("VEN-01", "Clientes", "Crear cliente nuevo con telefono y direccion.", "El cliente queda disponible para cotizaciones.", ""),
            ("VEN-02", "Clientes", "Editar datos de cliente.", "Los datos actualizados aparecen al seleccionar cliente.", ""),
            ("VEN-03", "Comercial", "Crear cotizacion de inventario disponible.", "La cotizacion queda en borrador/historial.", ""),
            ("VEN-04", "Comercial", "Crear preventa para cafe no disponible o en proceso.", "La preventa queda registrada para seguimiento.", ""),
            ("VEN-05", "Comercial", "Modificar precio negociado.", "El precio queda solo en esa cotizacion, no cambia precio base.", ""),
            ("VEN-06", "Comercial", "Cambiar cotizacion a enviada o aceptada.", "El estado se actualiza y queda visible para contabilidad.", ""),
            ("VEN-07", "Muestras", "Crear solicitud de muestra gratis.", "La muestra aparece como Gratis y Solicitada.", ""),
            ("VEN-08", "Muestras", "Crear solicitud de muestra con precio.", "La muestra muestra moneda y valor.", ""),
            ("VEN-09", "Muestras", "Revisar solicitudes propias.", "El vendedor ve seguimiento de sus muestras.", ""),
            ("VEN-10", "Documentos", "Abrir PDF/vista imprimible de cotizacion.", "Debe poder descargar o imprimir para enviar al cliente.", ""),
            ("VEN-11", "Ventas", "Ver estado de ventas asociadas.", "Debe ver estado general sin gestionar pagos internos.", ""),
        ],
    )

    add_use_cases(
        doc,
        "Muestras",
        "La persona encargada de muestras valida recepcion de solicitudes y avance de preparacion/entrega.",
        [
            ("MUE-01", "Muestras", "Entrar como usuario muestras.", "Debe abrir directamente el modulo Muestras.", ""),
            ("MUE-02", "Muestras", "Ver solicitudes creadas por vendedores/admin/contabilidad.", "La lista debe mostrar solicitante, telefono, cafe, cantidad, precio y fechas.", ""),
            ("MUE-03", "Muestras", "Cambiar solicitud a En preparacion.", "El estado cambia y se registra la gestion.", ""),
            ("MUE-04", "Muestras", "Cambiar solicitud a Lista.", "Queda claro que la muestra esta lista para entregar/enviar.", ""),
            ("MUE-05", "Muestras", "Cambiar solicitud a Entregada.", "La solicitud queda cerrada como entregada.", ""),
            ("MUE-06", "Muestras", "Cancelar una solicitud.", "El estado cambia a cancelada con confirmacion.", ""),
            ("MUE-07", "Muestras", "Agregar nota al cambiar estado.", "La nota queda visible para seguimiento.", ""),
            ("MUE-08", "Muestras", "Intentar crear solicitud desde usuario muestras.", "No debe aparecer formulario de creacion; solo gestion.", ""),
        ],
    )

    add_heading(doc, "Flujos completos recomendados", 1)
    add_matrix(
        doc,
        ["Flujo", "Pasos", "Resultado esperado"],
        [
            (
                "Cafe recibido a inventario",
                "Bodega recibe cafe -> Laboratorio aprueba -> Contabilidad registra compra/pago -> Inventario queda disponible.",
                "El lote pasa por estados correctos y conserva datos de trazabilidad.",
            ),
            (
                "Venta desde cotizacion",
                "Vendedor crea cotizacion -> la marca aceptada -> Contabilidad convierte a venta -> Bodega alista/despacha -> Contabilidad registra pagos.",
                "La venta queda trazable, descuenta inventario cuando corresponde y refleja saldo/pagos.",
            ),
            (
                "Proceso de cafe especial",
                "Bodega crea proceso con lotes disponibles -> Laboratorio finaliza proceso -> se crea lote PROC -> Contabilidad/Vendedor lo usan en venta.",
                "El sistema sabe que cafe entro al proceso y que lote procesado salio.",
            ),
            (
                "Solicitud de muestra",
                "Vendedor/admin/contabilidad crea muestra -> usuario muestras cambia estados -> se registra si fue gratis o cobrada.",
                "La persona encargada de muestras tiene lista clara de trabajo y seguimiento.",
            ),
        ],
        [1.25, 3.25, 2.0],
    )

    add_heading(doc, "Formato de feedback", 1)
    add_body(doc, "Registrar aqui los hallazgos durante las pruebas. Prioridad sugerida: Alta si bloquea el trabajo, Media si incomoda o genera confusion, Baja si es mejora visual o de comodidad.")
    add_matrix(
        doc,
        ["Rol", "Modulo", "Caso", "Problema o mejora", "Prioridad"],
        [("", "", "", "", "") for _ in range(10)],
        [1.0, 1.15, 0.75, 2.85, 0.75],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Casos de uso - Sistema Finca Anaya")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
